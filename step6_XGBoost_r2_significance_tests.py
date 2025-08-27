# r2_boxplot_and_wordtable.py
import os
import numpy as np
import matplotlib.pyplot as plt
import seaborn as sns
from scipy import stats

# =========================
# Paths & config
# =========================
data_root   = os.path.join(os.getcwd(), "eegfmri_data_07122025")
results_dir = os.path.join(os.getcwd(), "results")
fig_dir     = os.path.join(results_dir, "figs_r2")
os.makedirs(fig_dir, exist_ok=True)

TARGET_KEYS = ["Y_DAN", "Y_DMN", "Y_DNa", "Y_DNb"]
FILE_TMPL = "xgb_loso_r2_{key}_0to20s.npz"

# Save CSV/TSV?
WRITE_CSV = True
csv_table_path    = os.path.join(results_dir, "loso_inference_table_0to20s.tsv")
csv_long_r2_path  = os.path.join(results_dir, "loso_fold_r2_long_0to20s.tsv")

# Word table path (same folder as figure)
docx_path = os.path.join(fig_dir, "loso_inference_table_0to20s.docx")

# Figure style
sns.set(context="paper", style="whitegrid", font_scale=1.2)
FIGSIZE = (8.0, 4.6)
DPI = 300

# Names & colorblind-safe palette (Okabe–Ito)
NAME_MAP = {"Y_DAN": "DAN", "Y_DMN": "DMN", "Y_DNa": "DNa", "Y_DNb": "DNb"}
PALETTE  = {
    "Y_DAN": "#0072B2",  # blue
    "Y_DMN": "#E69F00",  # orange
    "Y_DNa": "#009E73",  # bluish green
    "Y_DNb": "#D55E00",  # vermillion
}

# =========================
# Helpers
# =========================
def ensure_dir(p):
    os.makedirs(p, exist_ok=True)

def sign_test_p(count_pos, n):
    # one-sided P(X >= count_pos) for X~Bin(n, 0.5)
    return float(np.sum(stats.binom.pmf(np.arange(count_pos, n + 1), n, 0.5)))

def ci95_t(mean, sd, n):
    se = sd / np.sqrt(n)
    tcrit = stats.t.ppf(0.975, df=n - 1)
    return (mean - tcrit * se, mean + tcrit * se)

def load_target_npz(key):
    path = os.path.join(results_dir, FILE_TMPL.format(key=key))
    if not os.path.exists(path):
        print(f"[WARN] Missing file for {key}: {path}")
        return None
    return np.load(path, allow_pickle=True)

def r2_stats(fold_r2):
    n = fold_r2.size
    mean = fold_r2.mean()
    sd   = fold_r2.std(ddof=1)
    ci   = ci95_t(mean, sd, n)
    median = float(np.median(fold_r2))
    positives = int((fold_r2 > 0).sum())

    # One-sided tests vs 0 (subjects as units)
    t_res = stats.ttest_1samp(fold_r2, 0.0, alternative="greater")
    try:
        w_stat, w_p = stats.wilcoxon(fold_r2, alternative="greater", zero_method="wilcox")
    except Exception:
        w_stat, w_p = (np.nan, np.nan)

    # Sign test (binomial)
    p_sign = sign_test_p(positives, n)

    # Effect size from Wilcoxon one-sided p (normal approx): z = ISF(p), dz=z/sqrt(n)
    if np.isfinite(w_p):
        z = stats.norm.isf(w_p)
        dz = float(z / np.sqrt(n))
    else:
        dz = np.nan

    return dict(n=n, mean=mean, sd=sd, ci=ci, median=median,
                positives=positives, t_stat=t_res.statistic, t_p=t_res.pvalue,
                w_V=w_stat, w_p=w_p, dz=dz, sign_p=p_sign)

def bh_fdr(pvals):
    """Benjamini–Hochberg FDR correction. Returns q-values."""
    p = np.asarray(pvals, dtype=float)
    n = p.size
    order = np.argsort(p)
    ranked = np.arange(1, n+1)
    p_sorted = p[order]
    q_sorted = np.minimum.accumulate((p_sorted * n / ranked)[::-1])[::-1]
    q = np.empty_like(p)
    q[order] = np.clip(q_sorted, 0, 1)
    return q

def common_ylim(all_values, pad=0.02):
    ymin = min(np.min(v) for v in all_values if v.size > 0)
    ymax = max(np.max(v) for v in all_values if v.size > 0)
    ymin = min(0.0, ymin)
    ymax = max(0.0, ymax)
    span = ymax - ymin
    return (ymin - pad*span, ymax + pad*span)

def savefig(fig, path_base):
    ensure_dir(os.path.dirname(path_base))
    for ext in (".png", ".pdf"):
        out = path_base + ext
        fig.savefig(out, dpi=DPI, bbox_inches="tight")
        print(f"Saved: {out}")
    plt.close(fig)

# =========================
# Load, compute stats, collect data
# =========================
rows = []
fold_r2_dict = {}
fold_subjects_dict = {}

for key in TARGET_KEYS:
    D = load_target_npz(key)
    if D is None:
        continue
    fold_r2 = np.array(D["fold_r2"], dtype=float)
    fold_subjects = D["fold_subjects"] if "fold_subjects" in D.files else np.array([f"{i+1:02d}" for i in range(len(fold_r2))])
    s = r2_stats(fold_r2)

    rows.append(dict(
        target_key=key, n=s["n"],
        mean_r2=s["mean"], sd_r2=s["sd"],
        ci_low=s["ci"][0], ci_high=s["ci"][1],
        median_r2=s["median"], positives=s["positives"],
        t_stat=s["t_stat"], t_p=s["t_p"],
        wilcoxon_V=s["w_V"], wilcoxon_p=s["w_p"],
        dz=s["dz"], sign_p=s["sign_p"]
    ))
    fold_r2_dict[key] = fold_r2
    fold_subjects_dict[key] = np.array(fold_subjects)

# FDR across targets (Wilcoxon p-values)
if rows:
    pvals = [r["wilcoxon_p"] for r in rows]
    qvals = bh_fdr(pvals)
    for r, q in zip(rows, qvals):
        r["q_fdr"] = float(q)

# =========================
# Print table (console) and write TSVs
# =========================
if rows:
    hdr = ("Target  N  MeanR2  [95% CI]        Median  +R2    Wilcoxon p  q(FDR)   dz")
    print(hdr)
    for r in rows:
        label = NAME_MAP.get(r['target_key'], r['target_key'])
        print(f"{label:6s} {r['n']:2d}  {r['mean_r2']:.3f}  "
              f"[{r['ci_low']:.3f},{r['ci_high']:.3f}]  "
              f"{r['median_r2']:.3f}  {r['positives']:2d}   "
              f"{r['wilcoxon_p']:.2e}  {r['q_fdr']:.2e}  {r['dz']:.2f}")

if WRITE_CSV and rows:
    cols = ["target_key","n","mean_r2","sd_r2","ci_low","ci_high","median_r2",
            "positives","t_stat","t_p","wilcoxon_V","wilcoxon_p","q_fdr","dz","sign_p"]
    with open(csv_table_path, "w") as f:
        f.write("\t".join(cols) + "\n")
        for r in rows:
            f.write("\t".join(str(r.get(c, "")) for c in cols) + "\n")
    print(f"\nSaved stats table: {csv_table_path}")

    with open(csv_long_r2_path, "w") as f:
        f.write("target_key\tsubject\tfold_r2\n")
        for key, vals in fold_r2_dict.items():
            subs = fold_subjects_dict[key]
            for s, v in zip(subs, vals):
                f.write(f"{NAME_MAP.get(key,key)}\t{s}\t{v:.6f}\n")
    print(f"Saved long-form fold R^2: {csv_long_r2_path}")

# =========================
# Single figure: BOX + strip (mean ± 95% CI overlay)
# (No q or dz annotations on the plot.)
# =========================
if not fold_r2_dict:
    raise SystemExit("No target files found; nothing to plot.")

# Prepare stacked data
stack_vals, stack_keys = [], []
present_keys = [k for k in TARGET_KEYS if k in fold_r2_dict]
order = [NAME_MAP.get(k, k) for k in present_keys]
palette = [PALETTE.get(k, "#333333") for k in present_keys]

for key in present_keys:
    vals = fold_r2_dict[key]
    stack_vals.extend(list(vals))
    stack_keys.extend([NAME_MAP.get(key, key)] * len(vals))

fig, ax = plt.subplots(figsize=FIGSIZE)

sns.boxplot(x=stack_keys, y=stack_vals, order=order, palette=palette,
            showfliers=False, width=0.6, linewidth=1.2, ax=ax)
sns.stripplot(x=stack_keys, y=stack_vals, order=order, color="#222222",
              size=3.6, alpha=0.65, jitter=0.15, ax=ax)

# overlay mean ± 95% CI per group as diamonds with error bars
def _common_ylim(all_values, pad=0.06):
    ymin = min(np.min(v) for v in all_values if v.size > 0)
    ymax = max(np.max(v) for v in all_values if v.size > 0)
    ymin = min(0.0, ymin)
    ymax = max(0.0, ymax)
    span = ymax - ymin
    return (ymin - pad*span, ymax + pad*span)

for i, key in enumerate(present_keys):
    vals = fold_r2_dict[key]
    n = len(vals)
    m = float(np.mean(vals))
    sd = float(np.std(vals, ddof=1))
    lo, hi = ci95_t(m, sd, n)
    ax.errorbar(i, m, yerr=[[m - lo], [hi - m]],
                fmt="D", color="white", mec="black", ms=6, capsize=4, lw=1.5, zorder=5)

ylims = _common_ylim(list(fold_r2_dict.values()), pad=0.06)
ax.axhline(0, color="#555555", lw=1, ls="--")
ax.set_xlabel("")
ax.set_ylabel(r"$R^2$")
ax.set_ylim(*ylims)
ax.set_title("LOSO performance by target (subject-level)")
sns.despine(fig=fig)

savefig(fig, os.path.join(fig_dir, "box_strip_r2_by_target"))

# =========================
# Save a Word table (.docx) with the stats (same folder as figure)
# =========================
try:
    from docx import Document
    from docx.shared import Pt, Inches

    ensure_dir(os.path.dirname(docx_path))
    doc = Document()
    doc.add_heading("LOSO R² by target (subject-level)", level=1)
    doc.add_paragraph(
        "Wilcoxon one-sided tests H₁: R² > 0; q-values are Benjamini–Hochberg FDR across targets. "
        "Effect size d_z = z/√n from the Wilcoxon normal approximation."
    )

    # Build table
    cols = ["Target","N","Mean R²","SD","95% CI","Median","Positives",
            "Wilcoxon p","q (FDR)","d\u2099"]  # dₙ (subscript n)
    table = doc.add_table(rows=1, cols=len(cols))
    table.style = "Light List Accent 1"
    hdr_cells = table.rows[0].cells
    for i, c in enumerate(cols):
        hdr_cells[i].text = c

    # Fill rows
    for r in rows:
        row = table.add_row().cells
        row[0].text = NAME_MAP.get(r['target_key'], r['target_key'])
        row[1].text = f"{int(r['n'])}"
        row[2].text = f"{r['mean_r2']:.3f}"
        row[3].text = f"{r['sd_r2']:.3f}"
        row[4].text = f"[{r['ci_low']:.3f}, {r['ci_high']:.3f}]"
        row[5].text = f"{r['median_r2']:.3f}"
        row[6].text = f"{int(r['positives'])}"
        row[7].text = f"{r['wilcoxon_p']:.2e}"
        row[8].text = f"{r.get('q_fdr', np.nan):.2e}"
        row[9].text = f"{r.get('dz', np.nan):.2f}"

    doc.save(docx_path)
    print(f"Saved Word table: {docx_path}")
except Exception as e:
    print(f"[WARN] Could not write Word table ({e}). "
          f'Install with: pip install python-docx. A TSV is available at: {csv_table_path}')
